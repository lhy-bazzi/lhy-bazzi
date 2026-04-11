"""DOCX parser — extracts text, headings, tables, images, and lists."""

from __future__ import annotations

import asyncio
import os
from typing import Any

from loguru import logger

from app.models.document import DocumentMetadata, ParsedDocument, DocumentElement
from app.models.enums import ElementType, FileType
from app.services.parsing.base import BaseParser
from app.utils.exceptions import ParseError


class DocxParser(BaseParser):
    """Microsoft Word (.docx) document parser."""

    def supported_types(self) -> list[FileType]:
        return [FileType.DOCX]

    async def parse(self, file_path: str, **kwargs: Any) -> ParsedDocument:
        doc_id = kwargs.get("doc_id", self._gen_id())
        filename = os.path.basename(file_path)
        logger.info("Parsing DOCX '{}'", filename)

        try:
            elements = await asyncio.to_thread(self._parse_sync, file_path)
        except Exception as exc:
            raise ParseError(f"DOCX parse failed: {exc}") from exc

        metadata = DocumentMetadata(file_size_bytes=self._file_size(file_path))
        return self._build_document(
            doc_id=doc_id,
            filename=filename,
            file_type="docx",
            elements=elements,
            parse_engine="python-docx",
            metadata=metadata,
        )

    def _parse_sync(self, file_path: str) -> list[DocumentElement]:
        from docx import Document as DocxDocument
        from docx.table import Table

        doc = DocxDocument(file_path)
        elements: list[DocumentElement] = []

        # --- Build a unified body iterator (paragraphs + tables in order) ---
        body = doc.element.body
        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                para = next(
                    (p for p in doc.paragraphs if p._element is child), None
                )
                if para is None:
                    continue
                text = para.text.strip()
                if not text:
                    continue

                style_name = (para.style.name or "").lower() if para.style else ""
                # Detect heading
                if style_name.startswith("heading"):
                    level = self._heading_level(style_name)
                    elements.append(self._build_element(
                        ElementType.HEADING, text, level=level
                    ))
                elif style_name.startswith("list"):
                    elements.append(self._build_element(ElementType.LIST, text))
                else:
                    elements.append(self._build_element(ElementType.TEXT, text))

            elif tag == "tbl":
                table = next(
                    (t for t in doc.tables if t._element is child), None
                )
                if table is not None:
                    md_table = self._table_to_markdown(table)
                    if md_table:
                        elements.append(self._build_element(ElementType.TABLE, md_table))

        return elements

    @staticmethod
    def _heading_level(style_name: str) -> int:
        for ch in style_name:
            if ch.isdigit():
                return int(ch)
        return 1

    @staticmethod
    def _table_to_markdown(table) -> str:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if len(rows) < 1:
            return ""
        # Insert separator after header row
        col_count = len(table.rows[0].cells) if table.rows else 0
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        return rows[0] + "\n" + separator + "\n" + "\n".join(rows[1:])
